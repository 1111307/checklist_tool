# -*- coding: utf-8 -*-
"""生成 配置核查作业指导书_v2.2.docx：
1. 提升第3、4章已有"补充核查方法"小节 Normal -> Heading 4
2. 修 4.23 编号错误（4.24.2 -> 4.23.2）
3. 移动 3.14.3（错位到第4章标题下）回到 3.14 条目
4. 删除错位正文（4.22.1 日志审计、4.33.1 备份）+ 4.13/4.32 占位
5. 插入 34 条"通用核查方法"补充小节
"""
from docx import Document
import re, json

SRC = '配置核查作业指导书_v2.1草稿_格式规范化.docx'
DST = '配置核查作业指导书_v2.2.docx'

doc = Document(SRC)

def chapter_bounds():
    b = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 2':
            m = re.match(r'(\d+)\s', p.text.strip())
            if m:
                b[int(m.group(1))] = i
    return b

cb = chapter_bounds()
ch3_start = cb.get(3, 0)
ch4_start = cb.get(4, 0)
ch5_start = cb.get(5, len(doc.paragraphs))
print(f'章节边界 ch3={ch3_start} ch4={ch4_start} ch5={ch5_start}')

# ---------- 阶段1：提升已有补充 Normal -> Heading4 ----------
promoted = 0
for i in range(ch3_start, ch5_start):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if p.style.name == 'Normal' and re.match(r'^\d+\.\d+\.\d+\s*(补充核查方法|通用核查方法)', t):
        p.style = doc.styles['Heading 4']
        promoted += 1
print(f'阶段1：提升 Heading4 {promoted} 条')

# ---------- 阶段2：修 4.23 编号 ----------
for i in range(ch4_start, ch5_start):
    p = doc.paragraphs[i]
    if p.style.name == 'Heading 4' and p.text.strip().startswith('4.24.2中标麒麟'):
        for r in p.runs:
            r.text = r.text.replace('4.24.2', '4.23.2')
print('阶段2：4.23 编号修复')

# ---------- 阶段3：移动 3.14.3 ----------
h2_ch4 = None
move_els = []
in_move = False
for i in range(ch3_start, ch5_start):
    p = doc.paragraphs[i]
    if p.style.name == 'Heading 2' and p.text.strip().startswith('4 '):
        h2_ch4 = p
        in_move = False
    elif h2_ch4 is not None and p.style.name == 'Heading 4' and p.text.strip().startswith('3.14.3'):
        in_move = True
        move_els.append(p._element)
        continue
    elif in_move:
        if p.style.name == 'Heading 3':
            break
        move_els.append(p._element)
if h2_ch4 is not None:
    for el in move_els:
        h2_ch4._element.addprevious(el)
print(f'阶段3：移动 3.14.3 共 {len(move_els)} 段')

# ---------- 阶段4：删除错位正文 ----------
def delete_normal_between(prefix_start, prefix_stop):
    paras = doc.paragraphs
    del_targets = []
    in_range = False
    for p in paras:
        t = p.text.strip()
        if p.style.name == 'Heading 4' and t.startswith(prefix_start):
            in_range = True
            continue
        if in_range and p.style.name == 'Heading 4':
            break
        if in_range and p.style.name == 'Normal':
            del_targets.append(p)
    for p in del_targets:
        p._element.getparent().remove(p._element)
    return len(del_targets)

n1 = delete_normal_between('4.22.1', '4.22.2')
n2 = delete_normal_between('4.33.1', '4.33.2')
print(f'阶段4：删除 4.22.1 错位 {n1} 段，4.33.1 错位 {n2} 段')

# 4.13/4.32 的"暂未找到资料"（仅这两条，第3、4章内）
n3 = 0
for i in range(ch3_start, ch5_start):
    p = doc.paragraphs[i]
    if p.style.name == 'Normal' and p.text.strip() == '暂未找到资料。':
        owner = None
        for j in range(i, ch3_start - 1, -1):
            if doc.paragraphs[j].style.name == 'Heading 3':
                m = re.match(r'(\d+\.\d+)', doc.paragraphs[j].text.strip())
                owner = m.group(1) if m else None
                break
        if owner in ('4.13', '4.32'):
            p._element.getparent().remove(p._element)
            n3 += 1
print(f'阶段4：删除 4.13/4.32 占位 {n3} 段')

# ---------- 阶段5：删除空 .1 标题 + 插入补充 ----------
def delete_empty_h4(h4_text):
    for p in list(doc.paragraphs):
        if p.style.name == 'Heading 4' and p.text.strip() == h4_text:
            p._element.getparent().remove(p._element)
            return True
    return False

# 4.22.2 标题去"（不全）"
for p in doc.paragraphs:
    if p.style.name == 'Heading 4' and '（不全）' in p.text:
        for r in p.runs:
            r.text = r.text.replace('（不全）', '')

delete_empty_h4('4.13.1Win7、WinXP')
delete_empty_h4('4.22.1Win7、WinXP')
delete_empty_h4('4.33.1Win7、WinXP')
print('阶段5a：删除空 .1 标题 3 个')

# 插入补充
existing = {'3.5', '3.11', '3.14', '4.1', '4.5', '4.8', '4.11', '4.19', '4.25', '4.29', '4.32'}
skip = existing | {'4.20', '4.24'}

ch3 = json.load(open('_final_ch3.json', encoding='utf-8'))
ch4 = json.load(open('_final_ch4.json', encoding='utf-8'))

def h3_map():
    h3 = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 3':
            m = re.match(r'(\d+\.\d+)\s', p.text.strip())
            if m:
                h3[m.group(1)] = i
    return h3

def next_h3_index(pid):
    h3 = h3_map()
    s = h3[pid]
    for v in sorted(h3.values()):
        if v > s:
            return v
    return len(doc.paragraphs)

def normalize_heading(it):
    h = it['heading']
    pid = it['id']
    if h.startswith(pid + ' 完整核查方法'):
        return h.replace(pid + ' 完整核查方法', pid + '.3 通用核查方法', 1)
    return h

def insert_before(anchor, style_name, text):
    p = anchor.insert_paragraph_before()
    p.style = doc.styles[style_name]
    if text:
        p.add_run(text)
    return p

def insert_blocks_before(anchor, blocks):
    # blocks: [(style_name, text), ...] 按文档顺序
    for style_name, text in reversed(blocks):
        insert_before(anchor, style_name, text)

items = ch3['items'] + ch4['items']
to_insert = [it for it in items if it['id'] not in skip]

inserted = 0
for it in to_insert:
    anchor_idx = next_h3_index(it['id'])
    anchor = doc.paragraphs[anchor_idx]
    heading = normalize_heading(it)
    blocks = [('Heading 4', heading)]
    for c in it['content']:
        blocks.append(('Normal', c))
    blocks.append(('Normal', it['judgment']))
    insert_blocks_before(anchor, blocks)
    inserted += 1

print(f'阶段5b：插入补充 {inserted} 条')

doc.save(DST)
print('已保存 ' + DST)
