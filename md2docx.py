# -*- coding: utf-8 -*-
"""Markdown -> docx 转换器（纯 python-docx，无需 pandoc）。

用法：python md2docx.py <输入.md> <输出.docx>
支持：标题(#/##/###)、表格、无序/有序列表、段落、**加粗**。
中文字体：正文宋体、标题黑体；标题黑色；表格全边框；正文首行缩进 2 字符。
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EAST_BODY = '宋体'
EAST_HEAD = '黑体'
LATIN = 'Times New Roman'

def _set_font(run, east, latin=LATIN, size=10.5, bold=False, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def _add_runs(p, text, east=EAST_BODY, size=10.5, bold=False):
    # 处理 **加粗**
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        b = part.startswith('**') and part.endswith('**')
        _set_font(p.add_run(part[2:-2] if b else part), east, size=size, bold=bold or b)

def _heading(doc, text, level):
    p = doc.add_heading(level=level)
    sizes = {1: 16, 2: 14, 3: 12}
    _add_runs(p, text, east=EAST_HEAD, size=sizes.get(level, 12), bold=True)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p

def _table(doc, rows):
    rows = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    t.autofit = True
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            p = c.paragraphs[0]
            _add_runs(p, cell, size=10.5, bold=(i == 0))
    return t

def convert(md_path, out_path):
    doc = Document()
    # 默认样式
    normal = doc.styles['Normal']
    normal.font.name = LATIN
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), EAST_BODY)
    normal.font.size = Pt(10.5)

    lines = open(md_path, encoding='utf-8').read().splitlines()
    i = 0
    n = len(lines)
    while i < n:
        ln = lines[i].rstrip()
        if ln.strip() == '':
            i += 1
            continue
        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)', ln)
        if m:
            _heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue
        # 表格
        if ln.lstrip().startswith('|'):
            tbl = [ln]
            while i + 1 < n and lines[i + 1].lstrip().startswith('|'):
                i += 1
                tbl.append(lines[i])
            # 去掉分隔行 |---|---|
            tbl = [r for r in tbl if not re.match(r'^\s*\|?[\s:|-]+\|?\s*$', r)]
            _table(doc, tbl)
            i += 1
            continue
        # 无序列表
        if re.match(r'^\s*[-*+]\s+', ln):
            p = doc.add_paragraph(style='List Bullet')
            _add_runs(p, re.sub(r'^\s*[-*+]\s+', '', ln))
            i += 1
            continue
        # 有序列表
        if re.match(r'^\s*\d+\.\s+', ln):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            _add_runs(p, ln.strip())
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph()
        _add_runs(p, ln)
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.5
        i += 1

    doc.save(out_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('用法：python md2docx.py <输入.md> <输出.docx>')
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
    print('已生成：' + sys.argv[2])
