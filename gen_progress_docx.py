# -*- coding: utf-8 -*-
"""生成《配置核查自动化工具 工作进展》Word 文档
内容：工作进展总结的一/二/三章节 + 附录《配置核查表 v2.0.0》矩阵表
依赖：python-docx、openpyxl
"""
import openpyxl
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"c:\Users\ryan.xiong\Desktop\配置核查"
XLSX = BASE + r"\配置核查表_v2.0.0.xlsx"
OUT = BASE + r"\配置核查自动化工具_工作进展.docx"

ZH_FONT = "微软雅黑"
TABLE_FONT = "宋体"


def set_run(run, size=11, bold=False, color=None, font=ZH_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def para(doc, text, size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13}
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=sizes.get(level, 12), bold=True, color=(0x1A, 0x3C, 0x6E))
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(doc, text, size=11):
    p = doc.add_paragraph()
    r = p.add_run("• " + text)
    set_run(r, size=size)
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def style_table(tbl):
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER


def fill_cell(cell, text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font=ZH_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, font=font)
    p.paragraph_format.space_after = Pt(2)


def simple_table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(tbl)
    for j, h in enumerate(headers):
        fill_cell(tbl.cell(0, j), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            fill_cell(tbl.cell(i, j), str(v))
    if widths:
        for j, w in enumerate(widths):
            for i in range(len(rows) + 1):
                tbl.cell(i, j).width = Mm(w)
    return tbl


doc = Document()
# 全局默认字体
st = doc.styles["Normal"]
st.font.name = ZH_FONT
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
# 页面 A4 纵向
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(22)

# ============ 标题 ============
p = doc.add_paragraph()
r = p.add_run("配置核查自动化工具 工作进展")
set_run(r, size=20, bold=True, color=(0x1A, 0x3C, 0x6E))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
p2 = doc.add_paragraph()
r2 = p2.add_run("工作主线：配置核查表 → 配置核查作业指导书 → 自动化核查工具")
set_run(r2, size=11, color=(0x60, 0x60, 0x60))
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(16)

# ============ 一、工作主线与依据文档 ============
heading(doc, "一、工作主线与依据文档")
para(doc, "以《配置核查表》为基准（检查项 × 检查对象矩阵），编制《配置核查作业指导书》，"
          "并据此开发自动化核查工具。三件套均已成型：")
simple_table(
    doc,
    ["产出物", "状态"],
    [
        ["配置核查表 v2.0.0：检查项 × 10类检查对象矩阵（Win7 / WinXP / 中标麒麟 / 银河麒麟 / "
         "Nginx / Tomcat / MySQL / SQLServer / 达梦 / Redis），打勾标识适用范围", "已定版"],
        ["配置核查作业指导书 v2.0.0：共10章（系统安全 / 用户安全 / 数据安全 / 应用安全 / 网络安全 / "
         "物理安全 / 组织机构 / 规章制度 / 管理实施 / 协议安全审计）", "已定版"],
        ["自动化核查工具（Windows 版 + 麒麟版）", "已可用，持续迭代中"],
    ],
    widths=[150, 25],
)

# ============ 二、自动化工具开发 ============
heading(doc, "二、自动化工具开发（核心产出）")
heading(doc, "1. Windows 版（win_xp7/）", level=2)
bullet(doc, "check_xp7.vbs + run.bat，覆盖 Windows XP / Windows 7 核查对象，137 项核查")

heading(doc, "2. 麒麟 Linux 版（kylin/）", level=2)
bullet(doc, "check_kylin.sh + run.sh，纯 Bash 实现、零依赖：目标机无需安装任何软件，"
            "拷入后 sudo bash run.sh 即可运行")
bullet(doc, "兼容银河麒麟（apt/dpkg 系）与中标麒麟（yum/rpm 系），自动识别发行版与包管理器")
bullet(doc, "137 项核查，与 Windows 版编号对齐，四类结论：")
simple_table(
    doc,
    ["结论", "含义", "数量级（真机实测）"],
    [
        ["合规 pass", "自动判定通过", "13 项"],
        ["不合规 fail", "自动判定发现问题", "9 项"],
        ["需人工核查 manual", "需结合管理制度/业务人工确认", "41 项"],
        ["不适用 na", "网络设备类、物理环境类等不适用于单机核查", "74 项"],
    ],
    widths=[35, 90, 50],
)
para(doc, "")
bullet(doc, "可自动判定的项包括：补丁安装、防火墙、SELinux/AppArmor、口令策略、PAM 复杂度、"
            "SSH 加固、USB 管控、审计服务、会话超时、高危服务、磁盘加密、日志保护等约 60 项")
bullet(doc, "报告双格式：HTML（浏览器查看）+ XLS（Excel/WPS 直接打开，无需安装 Office），"
            "按“未通过 → 需人工核查 → 不适用 → 通过”四分区展示")
bullet(doc, "每条结果附“参考指导书”列，标明对应指导书的章节及条款原文，方便测试人员现场翻阅对照")

# ============ 三、实测验证 ============
heading(doc, "三、实测验证")
para(doc, "已在真实银河麒麟 V10 SP1 机器上完整跑通，137 项全部出结果"
          "（13 合规 / 9 不合规 / 41 需人工核查 / 74 不适用）。"
          "检出的 9 项不合规均为真实安全问题，例如：")
for t in [
    "667 个软件包补丁未升级",
    "SELinux 处于 Disabled 状态",
    "密码有效期 99999 天（实际永不过期）",
    "未部署 USB 存储管控（usbguard / udev 均无）",
    "PAM 未配置口令复杂度策略",
    "/var/log 目录权限 775 过宽",
    "无登录会话超时（TMOUT / ClientAliveInterval 均未配置）",
    "SSH 仍使用默认 22 端口",
    "未禁止 root 直接远程登录（PermitRootLogin 未设置）",
]:
    bullet(doc, t)
para(doc, "结论：工具能真正发现安全隐患，不是形式化核查。")

# ============ 附录：配置核查表（横向页面） ============
land = doc.add_section(WD_SECTION.NEW_PAGE)
land.orientation = WD_ORIENT.LANDSCAPE
land.page_width, land.page_height = Mm(297), Mm(210)
land.left_margin = land.right_margin = Mm(12)
land.top_margin = land.bottom_margin = Mm(15)

heading(doc, "附录：配置核查表 v2.0.0（检查项 × 检查对象）")

wb = openpyxl.load_workbook(XLSX, data_only=True)
ws = wb["Sheet1"]

# 表头对象列（修正原表拼写：Tomocat→Tomcat, SQLSever→SQL Server）
obj_names = ["Win7", "WinXP", "中标麒麟", "银河麒麟", "Nginx", "Tomcat",
             "MySQL", "SQL Server", "达梦", "Redis", "其他"]
groups = [("操作系统", 2, 6), ("中间件", 6, 8), ("数据库", 8, 12), ("其他", 12, 13)]

# 数据行：第4行起，A列章节（合并单元格前向填充），B列检查项，C-M列标记
items = []
chapter = ""
for r in range(4, 138):
    a = ws.cell(r, 1).value
    if a is not None and str(a).strip():
        chapter = str(a).strip()
    item = ws.cell(r, 2).value
    if item is None or not str(item).strip():
        continue
    marks = []
    for c in range(3, 14):
        v = ws.cell(r, c).value
        v = "" if v is None else str(v).strip()
        marks.append(v)
    items.append((chapter, str(item).strip().replace("\n", " "), marks))

tbl = doc.add_table(rows=2 + len(items), cols=13)
style_table(tbl)

# ---- 两行表头 ----
HDR_BG = "1A3C6E"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


fill_cell(tbl.cell(0, 0), "章节", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(tbl.cell(0, 1), "检查项", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
tbl.cell(0, 0).merge(tbl.cell(1, 0))
tbl.cell(0, 1).merge(tbl.cell(1, 1))
for name, c1, c2 in groups:
    tbl.cell(0, c1).merge(tbl.cell(0, c2 - 1))
    fill_cell(tbl.cell(0, c1), name, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for j, n in enumerate(obj_names):
    fill_cell(tbl.cell(1, 2 + j), n, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for j in range(13):
    shade(tbl.cell(0, j), HDR_BG)
    shade(tbl.cell(1, j), HDR_BG)
    for rr in (0, 1):
        for pp in tbl.cell(rr, j).paragraphs:
            for rn in pp.runs:
                rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- 数据行 ----
row_of_item = {}
for i, (ch, item, marks) in enumerate(items):
    rr = 2 + i
    row_of_item[i] = rr
    fill_cell(tbl.cell(rr, 0), ch, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font=TABLE_FONT)
    fill_cell(tbl.cell(rr, 1), item, size=8, font=TABLE_FONT)
    for j, m in enumerate(marks):
        fill_cell(tbl.cell(rr, 2 + j), m, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font=TABLE_FONT)

# ---- 章节列纵向合并（与原表一致） ----
start = 0
for i in range(1, len(items) + 1):
    if i == len(items) or items[i][0] != items[start][0]:
        if i - 1 > start:
            tbl.cell(row_of_item[start], 0).merge(tbl.cell(row_of_item[i - 1], 0))
        start = i

# ---- 列宽 ----
widths = [14, 100, 11, 11, 13, 13, 12, 13, 12, 14, 12, 12, 11]
for j, w in enumerate(widths):
    for rr in range(2 + len(items)):
        tbl.cell(rr, j).width = Mm(w)

doc.save(OUT)
print("saved:", OUT)
print("items:", len(items))
chapters = {}
for ch, _, _ in items:
    chapters[ch] = chapters.get(ch, 0) + 1
for k, v in chapters.items():
    print("chapter:", k, v)
