# -*- coding: utf-8 -*-
"""v2.0.1 重建：从 v2.2 基线 + 插入 13 张终端核查截图。"""
import re
import os
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "配置核查作业指导书_v2.2.docx")
DST = os.path.join(BASE, "_v201_new.docx")
FIG = os.path.join(BASE, "_ch5_figs")

INSERTS = [
    ("5.3", "fig5_1v.png", 5.8, "图5-1 锐捷接入交换机 VLAN 划分与端口隔离核查示例"),
    ("5.4", "fig5_2v.png", 5.8, "图5-2 核心交换机路由表核查示例（display ip routing-table）"),
    ("5.5", "fig5_3v.png", 5.8, "图5-3 安全区域 VLAN 划分核查示例（display vlan）"),
    ("5.6", "fig5_4v.png", 5.8, "图5-4 非必要服务关闭核查示例（华为 VRP 命令回显）"),
    ("5.7", "fig5_5v.png", 5.8, "图5-5 唯一管理服务与加密通道核查示例（display ssh server status）"),
    ("5.8", "fig5_6v.png", 5.8, "图5-6 明文管理关闭与 SSH 加密核查示例"),
    ("5.10", "fig5_7v.png", 5.8, "图5-7 组播路由表核查示例（display multicast routing-table）"),
    ("5.11", "fig5_8v.png", 5.8, "图5-8 防火墙双机热备状态核查示例（display hrp state）"),
    ("5.12", "fig5_9v.png", 5.8, "图5-9 IPSec 隧道状态核查示例（display ike sa）"),
    ("5.16", "fig5_10v.png", 5.8, "图5-10 细粒度访问控制 ACL 核查示例（display acl all）"),
    ("5.18", "fig5_11v.png", 5.8, "图5-11 802.1x 接入认证核查示例（display dot1x）"),
    ("5.19", "fig5_12v.png", 5.8, "图5-12 接入层端口隔离组核查示例（display port-isolate group）"),
    ("5.23", "fig5_13v.png", 5.8, "图5-13 端口状态与设备配备必要性核查示例（display interface brief）"),
]


def make_caption(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '60'); spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman'); rf.set(qn('w:eastAsia'), '宋体')
    rPr.append(rf)
    rPr.append(OxmlElement('w:b'))
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '21'); rPr.append(sz)
    szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), '21'); rPr.append(szcs)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t)
    p.append(r)
    return p


def make_image_para(doc, img_path, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    el = p._p
    el.getparent().remove(el)
    return el


doc = Document(SRC)
body = doc.element.body

def get_lvl(el):
    pPr = el.pPr
    if pPr is None: return None
    ps = pPr.find(qn('w:pStyle'))
    if ps is None: return None
    return {'2':1,'3':2,'4':3,'5':4,'6':5}.get(ps.get(qn('w:val')))

def para_text(el):
    return ''.join(t.text or '' for t in el.findall('.//' + qn('w:t')))

# 定位第5章各节 H2 元素
h1_5 = None
sections = {}
for el in list(body):
    if el.tag != qn('w:p'):
        continue
    lv = get_lvl(el)
    t = para_text(el).strip()
    if lv == 1 and t.startswith('5 '):
        h1_5 = el
    elif lv == 1 and t.startswith('6 ') and h1_5 is not None:
        break
    elif lv == 2 and h1_5 is not None:
        m = re.match(r'^(5\.\d+)', t)
        if m:
            sections[m.group(1)] = el

def section_last_el(h2_el):
    node = h2_el.getnext()
    last = h2_el
    while node is not None:
        if node.tag == qn('w:p') and get_lvl(node) in (1, 2):
            break
        last = node
        node = node.getnext()
    return last

inserted = 0
for target, fname, width, caption in INSERTS:
    h2 = sections.get(target)
    if h2 is None:
        print(f'  [跳过] {target} 未找到')
        continue
    anchor = section_last_el(h2)
    img_el = make_image_para(doc, os.path.join(FIG, fname), width)
    cap_el = make_caption(caption)
    anchor.addnext(cap_el)
    anchor.addnext(img_el)
    inserted += 1
    print(f'  插入 {fname} → {target} 节末（{caption[:30]}…）')

print(f'共插入 {inserted} 张截图')
doc.save(DST)
print(f'已生成: {DST}')
