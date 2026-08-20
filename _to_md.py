# -*- coding: utf-8 -*-
"""将《配置核查作业指导书_v2.2.docx》转为 Markdown 副本（只读转换，不修改任何 docx）。"""
from docx import Document

SRC = r"C:\Users\ryan.xiong\Desktop\peizhitool\配置核查\配置核查作业指导书_v2.2.docx"
DST = r"C:\Users\ryan.xiong\Desktop\peizhitool\配置核查\配置核查作业指导书_v2.2.md"

doc = Document(SRC)
out = []
in_code = False

for p in doc.paragraphs:
    style = p.style.name
    text = p.text

    if style == 'HTML Preformatted':
        if not in_code:
            out.append('```')
            in_code = True
        out.append(text)
        continue
    else:
        if in_code:
            out.append('```')
            out.append('')
            in_code = False

    if style.startswith('Heading'):
        lvl = int(style.split()[-1])
        # 保证标题层级 >= 2（文档无 Heading 1）
        lvl = max(2, min(lvl, 6))
        out.append('#' * lvl + ' ' + text.strip())
    else:
        out.append(text)

if in_code:
    out.append('```')

with open(DST, 'w', encoding='utf-8') as f:
    f.write('\n\n'.join(out))

print(f"已生成: {DST}")
print(f"段落数: {len(doc.paragraphs)}")
print(f"输出字符数: {sum(len(x) for x in out)}")
