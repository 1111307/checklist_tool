# -*- coding: utf-8 -*-
"""第5章插入 8 张示意图 → 生成 配置核查作业指导书_v2.0.1.docx。"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:\Users\ryan.xiong\Desktop\peizhitool\配置核查\配置核查作业指导书_v2.2.docx"
DST = r"C:\Users\ryan.xiong\Desktop\peizhitool\配置核查\配置核查作业指导书_v2.0.1.docx"
FIG = r"C:\Users\ryan.xiong\Desktop\peizhitool\配置核查\_ch5_figs"

# (插图节号, 文件, 宽度in, 题注)
INSERTS = [
    ("chapter", "fig5_1.png", 2.4, "图5-1 网络安全核查总体框架"),
    ("5.1", "fig5_2.png", 2.6, "图5-2 跨网跨域数据交换审批与实施流程"),
    ("5.5", "fig5_3.png", 5.8, "图5-3 局域网安全区域划分示意"),
    ("5.8", "fig5_4.png", 4.5, "图5-4 唯一管理服务与加密管理通道示意"),
    ("5.15", "fig5_5.png", 5.8, "图5-5 网络边界等级隔离与边界防护示意"),
    ("5.18", "fig5_6.png", 5.4, "图5-6 802.1x 接入认证流程"),
    ("5.21", "fig5_7.png", 4.2, "图5-7 全网行为审计与攻击监测架构示意"),
    ("5.22", "fig5_8.png", 2.8, "图5-8 存储系统管理网络与应用网络分离示意"),
]


def make_caption(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '60'); spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman'); rf.set(qn('w:eastAsia'), '宋体')
    rPr.append(rf)
    rPr.append(OxmlElement('w:b'))
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '21'); rPr.append(sz)
    szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), '21'); rPr.append(szCs := OxmlElement('w:szCs'))
    szCs.set(qn('w:val'), '21')
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t)
    p.append(r)
    return p


def make_image_para(doc, img_path, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    # 从文末移动到目标位置
    el = p._p
    el.getparent().remove(el)
    return el


doc = Document(SRC)
body = doc.element.body

def get_lvl(el):
    pPr = el.pPr
    if pPr is None: return None
    ps = pPr.find(qn('w:pStyle'))
    if ps is None: return None
    return {'2':1,'3':2,'4':3,'5':4,'6':5}.get(ps.get(qn('w:val')))

def para_text(el):
    return ''.join(t.text or '' for t in el.findall('.//' + qn('w:t')))

# 定位第5章范围与各节末元素
from docx.text.paragraph import Paragraph
children = list(body)
h1_5 = None
sections = {}
for el in children:
    if el.tag != qn('w:p'):
        continue
    lv = get_lvl(el)
    t = para_text(el).strip()
    if lv == 1 and t.startswith('5 '):
        h1_5 = el
    elif lv == 1 and t.startswith('6 ') and h1_5 is not None:
        break
    elif lv == 2 and h1_5 is not None:
        m = re.match(r'^(5\.\d+)', t)
        if m:
            sections[m.group(1)] = el  # 记录 H2 元素（节首）

def section_last_el(h2_el):
    """从节首 H2 向后走到下一个 H1/H2 之前的最后一个元素。"""
    node = h2_el.getnext()
    last = h2_el
    while node is not None:
        if node.tag == qn('w:p') and get_lvl(node) in (1, 2):
            break
        last = node
        node = node.getnext()
    return last

inserted = 0
for target, fname, width, caption in INSERTS:
    if target == "chapter":
        anchor = h1_5  # 插在章标题后
    else:
        h2 = sections.get(target)
        if h2 is None:
            print(f'  [跳过] {target} 未找到')
            continue
        anchor = section_last_el(h2)
    img_el = make_image_para(doc, os.path.join(FIG, fname), width)
    cap_el = make_caption(caption)
    anchor.addnext(cap_el)
    anchor.addnext(img_el)
    inserted += 1
    print(f'  插入 {fname} → {target} 节末（{caption}）')

print(f'共插入 {inserted} 张图')
doc.save(DST)
print(f'已生成: {DST}')
